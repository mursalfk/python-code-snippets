import imaplib
import email
from email.header import decode_header
from email.utils import parsedate_to_datetime
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
from tkcalendar import DateEntry
import os
import re
import threading
from datetime import datetime
import zipfile
import csv

fetched_emails_data = []
date_filter_enabled = False
date_from_value = None
date_to_value = None
current_theme = "dark"

THEMES = {
    "dark": {
        "bg": "#1e1e1e",
        "fg": "#ffffff",
        "accent": "#007acc",
        "text": "#d4d4d4",
        "entry_bg": "#1e1e1e",
        "entry_fg": "#ffffff"
    },
    "light": {
        "bg": "#f0f0f0",
        "fg": "#000000",
        "accent": "#007acc",
        "text": "#000000",
        "entry_bg": "#ffffff",
        "entry_fg": "#000000"
    }
}

def sanitize_filename(name):
    name = re.sub(r'[\r\n\t]', ' ', name)
    name = re.sub(r'[\\/:"*?<>|]+', '', name)
    return name.strip()

def save_email_content(email_obj, base_path):
    msg = email_obj['msg']
    subject = email_obj['subject']
    date_str = email_obj['date']
    clean_subject = sanitize_filename(subject)[:50]
    folder_name = f"{clean_subject} [{date_str}]"
    folder_path = os.path.join(base_path, folder_name)
    os.makedirs(folder_path, exist_ok=True)

    doc = Document()
    doc.add_heading("Email Body", level=1)
    zipf = zipfile.ZipFile(os.path.join(folder_path, "attachments.zip"), 'w')

    if msg.is_multipart():
        for part in msg.walk():
            content_type = part.get_content_type()
            disposition = str(part.get("Content-Disposition"))
            if "attachment" not in disposition and content_type == "text/plain":
                body = part.get_payload(decode=True).decode(errors="ignore")
                doc.add_paragraph(body)
            elif "attachment" in disposition:
                filename = part.get_filename()
                if filename:
                    filename = sanitize_filename(filename)
                    filepath = os.path.join(folder_path, filename)
                    with open(filepath, "wb") as f:
                        payload = part.get_payload(decode=True)
                        f.write(payload)
                        zipf.write(filepath, arcname=filename)
                        os.remove(filepath)
    else:
        body = msg.get_payload(decode=True).decode(errors="ignore")
        doc.add_paragraph(body)

    doc.save(os.path.join(folder_path, "Body.docx"))
    zipf.close()

def fetch_emails_gui(email_addr, password, date_from, date_to, progress_bar, email_listbox, log_func):
    global fetched_emails_data
    fetched_emails_data = []
    try:
        log_func("📡 Connecting to mail server...")
        mail = imaplib.IMAP4_SSL("imap.gmail.com")
        log_func("🔐 Logging in...")
        mail.login(email_addr, password)
        log_func("📬 Selecting inbox...")
        mail.select("inbox")
        log_func("🔍 Searching for all emails...")
        status, messages = mail.search(None, "ALL")
        email_ids = messages[0].split()
        filtered = 0

        for idx, eid in enumerate(email_ids):
            status, msg_data = mail.fetch(eid, "(RFC822)")
            for response_part in msg_data:
                if isinstance(response_part, tuple):
                    msg = email.message_from_bytes(response_part[1])
                    subject, encoding = decode_header(msg["Subject"])[0]
                    if isinstance(subject, bytes):
                        subject = subject.decode(encoding if encoding else "utf-8")
                    subject = subject.strip() or "No_Subject"
                    date_obj = parsedate_to_datetime(msg["Date"])
                    date_str = date_obj.strftime("%Y-%m-%d")

                    if date_filter_enabled:
                        if date_from and date_obj < date_from:
                            continue
                        if date_to and date_obj > date_to:
                            continue

                    body = ""
                    if msg.is_multipart():
                        for part in msg.walk():
                            if part.get_content_type() == "text/plain":
                                body = part.get_payload(decode=True).decode(errors="ignore")
                                break
                    else:
                        body = msg.get_payload(decode=True).decode(errors="ignore")

                    fetched_emails_data.append({"msg": msg, "subject": subject, "date": date_str, "body": body})
                    email_listbox.insert(tk.END, f"{subject} [{date_str}]")
                    filtered += 1

            progress_bar["value"] = (idx + 1) * 100 / len(email_ids)
            progress_bar.update_idletasks()

        mail.logout()
        log_func(f"✅ Done. {filtered} emails loaded.")
    except Exception as e:
        log_func(f"❌ Error: {str(e)}")
        messagebox.showerror("Error", str(e))

def run_gui():
    def apply_theme():
        theme = THEMES[current_theme]
        root.configure(bg=theme["bg"])
        style.configure("TLabel", background=theme["bg"], foreground=theme["fg"])
        style.configure("TButton", background=theme["accent"], foreground=theme["fg"])
        style.configure("TEntry", fieldbackground=theme["entry_bg"], foreground=theme["entry_fg"])
        style.configure("Horizontal.TProgressbar", background=theme["accent"])
        for widget in all_widgets:
            try:
                widget.configure(bg=theme["bg"], fg=theme["fg"], insertbackground=theme["fg"])
            except:
                pass

    def toggle_theme():
        global current_theme
        current_theme = "light" if current_theme == "dark" else "dark"
        apply_theme()

    def browse_folder():
        folder = filedialog.askdirectory()
        if folder:
            save_path.set(folder)

    def start_fetching():
        email_val = email_entry.get()
        password_val = password_entry.get()
        email_listbox.delete(0, tk.END)
        status_box.delete(1.0, tk.END)
        preview_box.delete(1.0, tk.END)
        progress_bar["value"] = 0
        df = date_from_value if date_filter_enabled else None
        dt = date_to_value if date_filter_enabled else None
        threading.Thread(target=fetch_emails_gui, args=(
            email_val, password_val, df, dt,
            progress_bar, email_listbox, lambda msg: status_box.insert(tk.END, msg + "\n")
        ), daemon=True).start()

    def show_preview(event):
        preview_box.delete(1.0, tk.END)
        selection = email_listbox.curselection()
        if selection:
            body = fetched_emails_data[selection[0]]["body"]
            preview_box.insert(tk.END, body)

    def save_selected_emails():
        selected = email_listbox.curselection()
        if not selected:
            messagebox.showwarning("Nothing Selected", "Select emails first.")
            return
        for i in selected:
            save_email_content(fetched_emails_data[i], save_path.get())
        messagebox.showinfo("Success", "Selected emails saved.")

    def save_all_emails():
        for email_obj in fetched_emails_data:
            save_email_content(email_obj, save_path.get())
        messagebox.showinfo("Success", "All emails saved.")

    def export_to_csv():
        if not fetched_emails_data:
            messagebox.showwarning("No Data", "No emails to export.")
            return
        file = filedialog.asksaveasfilename(defaultextension=".csv", filetypes=[("CSV files", "*.csv")])
        if file:
            with open(file, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerow(["Subject", "Date"])
                for item in fetched_emails_data:
                    writer.writerow([item['subject'], item['date']])
            messagebox.showinfo("Exported", "Email list exported to CSV.")

    root = tk.Tk()
    root.title("📧 Email Fetcher")
    root.geometry("1100x650")

    style = ttk.Style()
    style.theme_use('clam')

    main_frame = tk.Frame(root)
    main_frame.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)

    left_frame = tk.Frame(main_frame)
    left_frame.grid(row=0, column=0, padx=10, sticky="nsew")

    right_frame = tk.Frame(main_frame)
    right_frame.grid(row=0, column=1, padx=10, sticky="nsew")

    main_frame.columnconfigure(0, weight=1)
    main_frame.columnconfigure(1, weight=1)

    email_entry = tk.Entry(left_frame, width=50)
    password_entry = tk.Entry(left_frame, show="*", width=50)
    save_path = tk.StringVar()
    save_entry = tk.Entry(left_frame, textvariable=save_path, width=38)
    browse_button = tk.Button(left_frame, text="Browse", command=browse_folder)
    date_option = ttk.Combobox(left_frame, values=["All", "Apply Date Filter"], state="readonly")
    date_option.set("All")
    fetch_button = tk.Button(left_frame, text="Fetch Emails", command=start_fetching)
    export_button = tk.Button(left_frame, text="Export to CSV", command=export_to_csv)
    theme_toggle = tk.Button(left_frame, text="Toggle Theme", command=toggle_theme)

    email_listbox = tk.Listbox(right_frame, selectmode=tk.MULTIPLE, width=60, height=10)
    save_selected_button = tk.Button(right_frame, text="Save Selected", command=save_selected_emails)
    save_all_button = tk.Button(right_frame, text="Save All", command=save_all_emails)
    preview_box = scrolledtext.ScrolledText(right_frame, width=60, height=8)
    status_box = scrolledtext.ScrolledText(right_frame, width=60, height=5)
    progress_bar = ttk.Progressbar(right_frame, length=500, mode='determinate')

    all_widgets = [root, main_frame, left_frame, right_frame, email_entry, password_entry, save_entry,
                   browse_button, fetch_button, export_button, theme_toggle, email_listbox, save_selected_button,
                   save_all_button, preview_box, status_box]

    # Layout
    for widget in [
        tk.Label(left_frame, text="Your Gmail:"), email_entry,
        tk.Label(left_frame, text="App Password:"), password_entry,
        tk.Label(left_frame, text="Save Folder:"),
        save_entry, browse_button,
        tk.Label(left_frame, text="Date Filter:"), date_option,
        fetch_button, export_button, theme_toggle
    ]:
        widget.pack(fill="x", pady=2)

    for widget in [
        tk.Label(right_frame, text="Fetched Emails:"), email_listbox,
        save_selected_button, save_all_button,
        tk.Label(right_frame, text="Email Preview:"), preview_box,
        tk.Label(right_frame, text="Status:"), status_box,
        tk.Label(right_frame, text="Progress:"), progress_bar
    ]:
        widget.pack(fill="x", pady=2)

    email_listbox.bind('<<ListboxSelect>>', show_preview)

    def on_date_option_change(event):
        global date_filter_enabled, date_from_value, date_to_value
        if date_option.get() == "Apply Date Filter":
            popup = tk.Toplevel(root)
            popup.title("Select Date Range")
            tk.Label(popup, text="From:").pack()
            df_picker = DateEntry(popup)
            df_picker.pack()
            tk.Label(popup, text="To:").pack()
            dt_picker = DateEntry(popup)
            dt_picker.pack()

            def apply_dates():
                global date_filter_enabled, date_from_value, date_to_value
                date_filter_enabled = True
                date_from_value = df_picker.get_date()
                date_to_value = dt_picker.get_date()
                popup.destroy()

            tk.Button(popup, text="Apply", command=apply_dates).pack()

        else:
            date_filter_enabled = False

    date_option.bind("<<ComboboxSelected>>", on_date_option_change)
    apply_theme()
    root.mainloop()

if __name__ == "__main__":
    run_gui()
